"""Convert bug4_report.html to bug4_report.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import re, sys

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get(side, 'none'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2e3248')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── document setup ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── styles ──────────────────────────────────────────────────────────────────

styles = doc.styles

def _style(name, base, size, bold=False, color=None, space_before=0, space_after=0):
    try:
        s = styles[name]
    except KeyError:
        s = styles.add_style(name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    s.base_style = styles[base]
    f = s.font
    f.size = Pt(size)
    f.bold = bold
    if color:
        r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        f.color.rgb = RGBColor(r, g, b)
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after  = Pt(space_after)
    return s

_style('MyTitle',    'Normal', 20, bold=True,  color='4f46e5', space_before=0, space_after=6)
_style('MySubtitle', 'Normal', 10, bold=False, color='768390', space_before=0, space_after=14)
_style('MyH2',       'Normal', 14, bold=True,  color='e2e8f0', space_before=16, space_after=6)
_style('MyH3',       'Normal', 11, bold=True,  color='56cfb2', space_before=10, space_after=4)
_style('MyH4',       'Normal', 10, bold=True,  color='adbac7', space_before=8,  space_after=3)
_style('MyBody',     'Normal', 10, bold=False, color='cdd9e5', space_before=0,  space_after=6)
_style('MyCode',     'Normal',  9, bold=False, color='adbac7', space_before=0,  space_after=0)
_style('MyCaption',  'Normal',  8, bold=False, color='768390', space_before=0,  space_after=4)
_style('MyBullet',   'Normal', 10, bold=False, color='cdd9e5', space_before=0,  space_after=3)

# Set document background to dark
body = doc.element.body
sectPr = body.get_or_add_sectPr() if hasattr(body, 'get_or_add_sectPr') else None
bg = OxmlElement('w:background')
bg.set(qn('w:color'), '0F1117')
doc.element.insert(0, bg)
doc.settings.element.append(OxmlElement('w:displayBackgroundShape'))

# ── parse HTML ──────────────────────────────────────────────────────────────

with open('bug4_report.html', encoding='utf-8') as f:
    soup = BeautifulSoup(f.read(), 'html.parser')

def clean(text):
    return re.sub(r'\s+', ' ', text).strip()

def add_code_block(doc, code_text):
    """Add a shaded code block paragraph by paragraph."""
    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph(style='MyCode')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xad, 0xba, 0xc7)
        # light shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '161B22')
        pPr.append(shd)

def add_math_block(doc, text):
    lines = [l for l in text.split('\n') if l.strip()]
    for line in lines:
        p = doc.add_paragraph(style='MyCode')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7c, 0xc4, 0xfa)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '161B22')
        pPr.append(shd)

# ── TITLE BLOCK ──────────────────────────────────────────────────────────────

hero = soup.find(class_='hero')
if hero:
    h1 = hero.find('h1')
    if h1:
        p = doc.add_paragraph(clean(h1.get_text()), style='MyTitle')
    sub = hero.find(class_='subtitle')
    if sub:
        p = doc.add_paragraph(clean(sub.get_text()), style='MySubtitle')

add_horizontal_rule(doc)

# ── SECTIONS ─────────────────────────────────────────────────────────────────

main = soup.find('main')
if not main:
    main = soup.body

for section in main.find_all('section', recursive=False):
    for child in section.children:
        if not hasattr(child, 'name') or child.name is None:
            continue
        tag = child.name.lower()
        text = clean(child.get_text())

        if tag == 'h2':
            add_horizontal_rule(doc)
            p = doc.add_paragraph(text, style='MyH2')

        elif tag == 'h3':
            p = doc.add_paragraph(text, style='MyH3')

        elif tag == 'h4':
            p = doc.add_paragraph(text, style='MyH4')

        elif tag in ('p',):
            if text:
                p = doc.add_paragraph(text, style='MyBody')

        elif tag == 'pre':
            code_text = child.get_text()
            add_code_block(doc, code_text)
            doc.add_paragraph('', style='MyCaption').paragraph_format.space_after = Pt(4)

        elif tag == 'div' and 'math-block' in child.get('class', []):
            txt = child.get_text('\n')
            add_math_block(doc, txt)
            doc.add_paragraph('', style='MyCaption').paragraph_format.space_after = Pt(4)

        elif tag == 'div' and ('card' in child.get('class', []) or 'callout' in child.get('class', [])):
            # render card/callout as an indented block
            for sub in child.children:
                if not hasattr(sub, 'name') or sub.name is None:
                    continue
                stag = sub.name.lower()
                stxt = clean(sub.get_text())
                if stag in ('h3', 'h4'):
                    p = doc.add_paragraph(stxt, style='MyH3')
                    p.paragraph_format.left_indent = Cm(0.5)
                elif stag == 'p' and stxt:
                    p = doc.add_paragraph(stxt, style='MyBody')
                    p.paragraph_format.left_indent = Cm(0.5)
                elif stag == 'pre':
                    add_code_block(doc, sub.get_text())
                elif stag == 'div' and 'math-block' in sub.get('class', []):
                    add_math_block(doc, sub.get_text('\n'))
                elif stag == 'div' and 'pros-cons' in sub.get('class', []):
                    for ul in sub.find_all('ul'):
                        cls = ul.get('class', [])
                        is_pro = 'pros' in cls
                        for li in ul.find_all('li'):
                            prefix = '+ ' if is_pro else '− '
                            col = RGBColor(0x56,0xcf,0xb2) if is_pro else RGBColor(0xf4,0x70,0x67)
                            p = doc.add_paragraph(style='MyBullet')
                            p.paragraph_format.left_indent = Cm(1)
                            run = p.add_run(prefix)
                            run.font.bold = True
                            run.font.color.rgb = col
                            run2 = p.add_run(clean(li.get_text()))
                            run2.font.size = Pt(9.5)
                elif stag == 'ul':
                    for li in sub.find_all('li'):
                        p = doc.add_paragraph(style='MyBullet')
                        p.paragraph_format.left_indent = Cm(0.8)
                        run = p.add_run('• ' + clean(li.get_text()))

        elif tag == 'div':
            # solution-grid: iterate sol-cards
            for sol in child.find_all(class_='sol-card'):
                hdr = sol.find(class_='sol-header')
                if hdr:
                    num_el = hdr.find(class_='sol-num')
                    title_el = hdr.find(class_='sol-title')
                    tag_el = hdr.find(class_='sol-tag')
                    title_txt = ''
                    if num_el: title_txt += f"[{clean(num_el.get_text())}] "
                    if title_el: title_txt += clean(title_el.get_text())
                    if tag_el: title_txt += f"  [{clean(tag_el.get_text())}]"
                    p = doc.add_paragraph(title_txt, style='MyH3')
                for sub in sol.children:
                    if not hasattr(sub, 'name') or sub.name is None:
                        continue
                    stag = sub.name.lower()
                    stxt = clean(sub.get_text())
                    if stag == 'p' and stxt:
                        p = doc.add_paragraph(stxt, style='MyBody')
                        p.paragraph_format.left_indent = Cm(0.3)
                    elif stag == 'div' and 'complexity' in sub.get('class', []):
                        p = doc.add_paragraph(stxt, style='MyCaption')
                        p.paragraph_format.left_indent = Cm(0.3)
                    elif stag == 'pre':
                        add_code_block(doc, sub.get_text())
                    elif stag == 'div' and 'pros-cons' in sub.get('class', []):
                        for ul in sub.find_all('ul'):
                            cls = ul.get('class', [])
                            is_pro = 'pros' in cls
                            for li in ul.find_all('li'):
                                prefix = '+ ' if is_pro else '− '
                                col = RGBColor(0x56,0xcf,0xb2) if is_pro else RGBColor(0xf4,0x70,0x67)
                                p = doc.add_paragraph(style='MyBullet')
                                p.paragraph_format.left_indent = Cm(0.8)
                                run = p.add_run(prefix)
                                run.font.bold = True
                                run.font.color.rgb = col
                                run2 = p.add_run(clean(li.get_text()))
                                run2.font.size = Pt(9.5)
                doc.add_paragraph('', style='MyCaption')

        elif tag == 'table':
            headers = [clean(th.get_text()) for th in child.find('thead').find_all('th')] if child.find('thead') else []
            rows = []
            tbody = child.find('tbody')
            if tbody:
                for tr in tbody.find_all('tr'):
                    rows.append([clean(td.get_text()) for td in tr.find_all('td')])
            if not headers and rows:
                headers = rows.pop(0)
            if headers:
                ncols = len(headers)
                tbl = doc.add_table(rows=1+len(rows), cols=ncols)
                tbl.style = 'Table Grid'
                # header row
                hrow = tbl.rows[0]
                for idx, hdr_txt in enumerate(headers):
                    cell = hrow.cells[idx]
                    cell.text = hdr_txt
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xad,0xba,0xc7)
                    set_cell_bg(cell, '242736')
                # data rows
                for ri, row_data in enumerate(rows):
                    trow = tbl.rows[ri+1]
                    for ci, cell_txt in enumerate(row_data):
                        cell = trow.cells[ci]
                        cell.text = cell_txt
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xcd,0xd9,0xe5)
                        set_cell_bg(cell, '1A1D27')
                doc.add_paragraph('')

# ── footer line ──────────────────────────────────────────────────────────────
add_horizontal_rule(doc)
p = doc.add_paragraph('Generated 2026-03-19 — MIE8888 WSN Cooperative Tracking — Bug 4 Analysis',
                      style='MyCaption')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── save ──────────────────────────────────────────────────────────────────────
out = 'bug4_report.docx'
doc.save(out)
print(f'Saved: {out}')
