"""Convert bug4_report_v2.html → bug4_report_v2.docx  (light/professional theme)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString
import re, copy

# ── palette ────────────────────────────────────────────────────────────────
C = dict(
    navy   = RGBColor(0x1e,0x29,0x3b),
    accent = RGBColor(0x25,0x63,0xeb),
    teal   = RGBColor(0x0d,0x94,0x88),
    amber  = RGBColor(0xd9,0x77,0x06),
    red    = RGBColor(0xdc,0x26,0x26),
    body   = RGBColor(0x1e,0x29,0x3b),
    muted  = RGBColor(0x64,0x74,0x8b),
    white  = RGBColor(0xff,0xff,0xff),
    light  = RGBColor(0xf1,0xf5,0xf9),
    green  = RGBColor(0x06,0x5f,0x46),
)
HEX = dict(navy='1E293B', accent='2563EB', teal='0D9488',
           amber='D97706', light='F8FAFC', border='E2E8F0',
           code_bg='F8FAFC', alg_header='1E293B', white='FFFFFF',
           stripe='F8FAFC')

def rgb_hex(h): r,g,b=int(h[0:2],16),int(h[2:4],16),int(h[4:6],16); return RGBColor(r,g,b)

# ── xml helpers ────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def para_bg(p, hex_color):
    pPr=p._p.get_or_add_pPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); pPr.append(shd)

def para_border_left(p, hex_color, sz='24'):
    pPr=p._p.get_or_add_pPr()
    pBdr=OxmlElement('w:pBdr')
    left=OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),sz)
    left.set(qn('w:space'),'4'); left.set(qn('w:color'),hex_color)
    pBdr.append(left); pPr.append(pBdr)

def h_rule(doc):
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr()
    pBdr=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),HEX['accent'])
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    return p

def clean(t): return re.sub(r'\s+',' ',t).strip()

# ── document ───────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2)
    s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

# default style
doc.styles['Normal'].font.name='Calibri'
doc.styles['Normal'].font.size=Pt(10.5)
doc.styles['Normal'].font.color.rgb=C['body']

# ── style factory ──────────────────────────────────────────────────────────
def mkstyle(name,size,bold=False,color=None,sb=0,sa=4,italic=False):
    try: st=doc.styles[name]
    except KeyError: st=doc.styles.add_style(name,1)
    st.base_style=doc.styles['Normal']
    f=st.font; f.size=Pt(size); f.bold=bold; f.italic=italic
    if color: f.color.rgb=color
    pf=st.paragraph_format; pf.space_before=Pt(sb); pf.space_after=Pt(sa)
    return st

mkstyle('RP_Title',    20, True,  C['navy'],   sb=0,  sa=6)
mkstyle('RP_Sub',      10, False, C['muted'],  sb=0,  sa=14)
mkstyle('RP_H2',       14, True,  C['navy'],   sb=14, sa=6)
mkstyle('RP_H3',       11, True,  C['teal'],   sb=10, sa=4)
mkstyle('RP_H4',       10, True,  C['body'],   sb=8,  sa=3)
mkstyle('RP_Body',     10.5,False,C['body'],   sb=0,  sa=6)
mkstyle('RP_Code',      8.5,False,C['muted'],  sb=0,  sa=0)
mkstyle('RP_Math',      9,  False,C['accent'], sb=0,  sa=0)
mkstyle('RP_Caption',   8,  False,C['muted'],  sb=2,  sa=6)
mkstyle('RP_Bullet',   10,  False,C['body'],   sb=0,  sa=3)

# ── helpers ────────────────────────────────────────────────────────────────
def add_code(doc, text, math=False):
    bg = HEX['code_bg']
    col = C['accent'] if math else C['body']
    fname = 'Consolas'
    for line in text.split('\n'):
        p=doc.add_paragraph(style='RP_Code')
        p.paragraph_format.left_indent=Cm(0.5)
        p.paragraph_format.space_before=Pt(0)
        p.paragraph_format.space_after=Pt(0)
        run=p.add_run(line if line else ' ')
        run.font.name=fname; run.font.size=Pt(8.5 if not math else 9)
        run.font.color.rgb=col
        para_bg(p, bg)

def add_callout(doc, text, border_hex):
    p=doc.add_paragraph(style='RP_Body')
    p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(4)
    p.paragraph_format.space_after=Pt(4)
    p.add_run(text)
    para_bg(p, HEX['light'])
    para_border_left(p, border_hex)

def add_cover_row(doc, label, value):
    p=doc.add_paragraph(style='RP_Body')
    r1=p.add_run(label+': '); r1.font.bold=True; r1.font.color.rgb=C['muted']
    r2=p.add_run(value);      r2.font.color.rgb=C['light']

def add_section_h2(doc, text):
    h_rule(doc)
    p=doc.add_paragraph(text, style='RP_H2')
    return p

def add_table(doc, headers, rows):
    ncols=len(headers)
    tbl=doc.add_table(rows=1+len(rows), cols=ncols)
    tbl.style='Table Grid'
    # header
    hr=tbl.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h
        c.paragraphs[0].runs[0].font.bold=True
        c.paragraphs[0].runs[0].font.size=Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb=C['white']
        cell_bg(c, HEX['alg_header'])
    # rows
    for ri,row in enumerate(rows):
        tr=tbl.rows[ri+1]
        bg=HEX['stripe'] if ri%2==1 else HEX['white']
        for ci,val in enumerate(row):
            c=tr.cells[ci]; c.text=val
            c.paragraphs[0].runs[0].font.size=Pt(9)
            cell_bg(c, bg)
    doc.add_paragraph('',style='RP_Caption')

# ── parse ──────────────────────────────────────────────────────────────────
with open('bug4_report_v2.html', encoding='utf-8') as f:
    soup=BeautifulSoup(f.read(),'html.parser')

# ══ COVER ══
cover=soup.find(class_='cover')
if cover:
    p=doc.add_paragraph(style='RP_Body')
    r=p.add_run('MIE8888 — Graduate Project Report')
    r.font.size=Pt(8.5); r.font.color.rgb=C['muted']; r.font.bold=True

    h1=cover.find('h1')
    if h1:
        p=doc.add_paragraph(clean(h1.get_text()), style='RP_Title')

    sub=cover.find(class_='subtitle')
    if sub:
        p=doc.add_paragraph(clean(sub.get_text()), style='RP_Sub')

    meta=cover.find(class_='cover-meta')
    if meta:
        for div in meta.find_all('div'):
            strong=div.find('strong')
            if strong:
                label=clean(strong.get_text())
                strong.extract()
                val=clean(div.get_text())
                add_cover_row(doc, label, val)

    doc.add_paragraph('', style='RP_Body')

# ══ ABSTRACT ══
abstract=soup.find(class_='abstract')
if abstract:
    text=clean(abstract.get_text())
    add_callout(doc, text, HEX['accent'])
    doc.add_paragraph('', style='RP_Caption')

# ══ SECTIONS ══
main=soup.find('main')
for section in (main or soup.body).find_all('section', recursive=False):
    for child in section.children:
        if isinstance(child, NavigableString): continue
        tag=child.name.lower()
        text=clean(child.get_text())
        cls=child.get('class',[])

        if tag=='h2':
            add_section_h2(doc, text)

        elif tag=='h3':
            doc.add_paragraph(text, style='RP_H3')

        elif tag=='h4':
            doc.add_paragraph(text, style='RP_H4')

        elif tag=='p' and text:
            doc.add_paragraph(text, style='RP_Body')

        elif tag=='pre':
            add_code(doc, child.get_text())
            doc.add_paragraph('', style='RP_Caption')

        elif tag=='div' and 'math-block' in cls:
            add_code(doc, child.get_text('\n'), math=True)
            doc.add_paragraph('', style='RP_Caption')

        elif tag=='div' and 'callout' in ' '.join(cls):
            if 'callout-warn'  in cls: border=HEX['amber']
            elif 'callout-green' in cls: border=HEX['teal']
            else: border=HEX['accent']
            add_callout(doc, text, border)

        elif tag=='table':
            thead=child.find('thead')
            headers=[clean(th.get_text()) for th in thead.find_all('th')] if thead else []
            rows=[]
            tbody=child.find('tbody')
            if tbody:
                for tr in tbody.find_all('tr'):
                    rows.append([clean(td.get_text()) for td in tr.find_all('td')])
            if headers:
                add_table(doc, headers, rows)

        elif tag=='div' and 'alg-card' in cls:
            # header
            hdr=child.find(class_='alg-header')
            if hdr:
                num=hdr.find(class_='alg-num')
                title=hdr.find(class_='alg-title')
                badge=hdr.find(class_='alg-badge')
                parts=[]
                if num: parts.append(f"[{clean(num.get_text())}]")
                if title: parts.append(clean(title.get_text()))
                if badge: parts.append(f"({clean(badge.get_text())})")
                p=doc.add_paragraph(' '.join(parts), style='RP_H3')
                p.paragraph_format.space_before=Pt(14)

            body_div=child.find(class_='alg-body')
            if body_div:
                for el in body_div.children:
                    if isinstance(el, NavigableString): continue
                    etag=el.name.lower()
                    etext=clean(el.get_text())
                    ecls=el.get('class',[])

                    if etag=='p' and etext:
                        p=doc.add_paragraph(etext, style='RP_Body')
                        p.paragraph_format.left_indent=Cm(0.3)

                    elif etag=='div' and 'complexity' in ecls:
                        p=doc.add_paragraph(etext, style='RP_Caption')
                        p.paragraph_format.left_indent=Cm(0.3)
                        para_bg(p, HEX['light'])

                    elif etag=='pre':
                        add_code(doc, el.get_text())
                        doc.add_paragraph('', style='RP_Caption')

                    elif etag=='div' and 'props' in ecls:
                        for col in el.find_all(class_='props-col'):
                            h5=col.find('h5')
                            if h5:
                                is_pro='pro' in h5.get('class',[])
                                color=C['teal'] if is_pro else C['red']
                                ph=doc.add_paragraph(style='RP_Bullet')
                                ph.paragraph_format.left_indent=Cm(0.5)
                                rh=ph.add_run(clean(h5.get_text()))
                                rh.font.bold=True; rh.font.color.rgb=color
                            for li in col.find_all('li'):
                                p=doc.add_paragraph(style='RP_Bullet')
                                p.paragraph_format.left_indent=Cm(0.8)
                                li_text=clean(li.get_text())
                                r=p.add_run(li_text)
                                r.font.size=Pt(9.5)

            doc.add_paragraph('', style='RP_Caption')

# ══ FOOTER ══
h_rule(doc)
p=doc.add_paragraph(
    'MIE8888 Graduate Project — Multi-Target Cooperative Tracking in Mobile WSN — March 2026',
    style='RP_Caption')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER

doc.save('bug4_report_v2.docx')
print('Saved: bug4_report_v2.docx')
