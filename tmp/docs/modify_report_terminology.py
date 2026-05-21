from __future__ import annotations

from dataclasses import dataclass
from difflib import SequenceMatcher
from html import escape
from pathlib import Path
import re

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE_DOCX = Path(r"C:\Users\abel\Desktop\Yeqi-mie8888-final-report.docx")
OUTPUT_DIR = Path(r"C:\Users\abel\Desktop\MIE8888\mie8888\output\doc")
OUTPUT_DOCX = OUTPUT_DIR / "Yeqi-mie8888-final-report_expanded-joint-assignment.docx"
OUTPUT_HTML = OUTPUT_DIR / "Yeqi-mie8888-final-report_expanded-joint-assignment_highlighted.html"


PARA_UPDATES = {
    0: "An Expanded Joint Assignment Approach for Cooperative Multi-Target Tracking in Mobile WSNs",
    6: (
        'In this report, the proposed method is referred to as an expanded joint assignment approach. '
        'Here, "expanded" means that the joint assignment scope includes not only the current calling targets, '
        'but also any previously served target whose committed interceptor is contested by a later handover request. '
        'The proposed expanded joint assignment approach differs from the previous approach in three ways. '
        'First, it includes both the current calling targets and any previously served target affected by a contested '
        'committed interceptor in the same joint assignment problem. Second, it adopts a two-stage assignment rule '
        'that first maximizes the number of filled interceptor slots and then minimizes the total assignment cost '
        'among assignments with the same number of filled slots. Third, it replaces non-overlapping two-sensor '
        'team comparison with an event-local optimal sensor-to-target assignment over all involved targets and all '
        'eligible candidate sensors. Under ideal communication assumptions in simulation, the resulting formulation '
        'retains the distributed execution logic of the original DES framework while providing a different mechanism '
        'for resolving interceptor assignment conflicts; the communication requirements and implications for '
        'physical-platform implementation are discussed later in the report.'
    ),
    11: (
        'I formulate committed-interceptor conflict resolution as an expanded joint assignment problem, where the '
        'assignment scope includes both current calling targets and previously served targets affected by contested '
        'committed interceptors. Compared with the previous MinMax-based method, the proposed expanded joint assignment '
        'approach differs in three ways. First, it includes both the current calling targets and any previously served '
        'target affected by a contested committed interceptor in the same joint assignment problem. Second, it adopts '
        'a two-stage assignment rule that first maximizes the number of filled interceptor slots and then minimizes '
        'the total assignment cost among assignments with the same number of filled slots. Third, it replaces the '
        'previous comparison of non-overlapping two-sensor teams with a constrained sensor-to-target assignment model '
        'that jointly enforces the interceptor capacity of each target and the one-target-at-a-time constraint of each '
        'sensor. Under ideal communication assumptions in simulation, this revised formulation retains the distributed '
        'execution logic of the original DES framework while providing a different mechanism for resolving interceptor '
        'assignment conflicts. The communication requirements and implications for physical-platform implementation are '
        'discussed later in the report.'
    ),
    18: (
        'This report builds upon [1] by formulating committed-interceptor conflict resolution as an expanded joint '
        'assignment problem. When a committed interceptor is contested by a later handover request, the target that '
        'may lose that interceptor is included in the expanded joint assignment before the assignment problem is '
        'solved. The report then applies an event-local two-stage assignment rule over the involved targets and '
        'eligible candidate sensors. Unlike the MinMax objective, which emphasizes worst-team-cost fairness, the '
        'proposed two-stage assignment rule first maximizes the number of filled interceptor slots and then minimizes '
        'the total assignment cost among feasible assignments with the same number of filled slots.'
    ),
    26: (
        'The system assumptions used in the remainder of the report are as follows: the workspace is planar, sensors '
        'are homogeneous, assignments are event-triggered, and each sensor can serve at most one target at a time '
        'within the current assignment event. Because different targets may enter the assignment process at different '
        'times, a later handover request may compete with sensors that have already been committed to another target. '
        'The proposed method addresses this case by expanding the joint assignment scope before solving the assignment '
        'problem.'
    ),
    29: (
        'When one or more targets issue handover requests, the system defines the targets included in the expanded '
        'joint assignment and a set of candidate sensors. The targets included in the expanded joint assignment include '
        'the current calling targets and any previously served targets whose committed interceptors may be reassigned '
        'due to the current conflict-resolution event.'
    ),
    51: "Conflict Detection and Expanded Joint Assignment Scope",
    56: (
        'To address this issue, the proposed expanded joint assignment approach enlarges the joint assignment scope '
        'when necessary. The joint assignment initially includes only the current calling targets in the assignment '
        'event. If the conflict involves only these current handover requests, no further scope expansion is needed. '
        'The scope is expanded only when a committed-interceptor overlap is detected. In that case, the target that '
        'may lose its committed interceptor is added to the joint assignment. The resulting expanded joint assignment '
        'scope is then solved once so that any target affected by the conflict is considered within the same assignment '
        'step.'
    ),
    58: "Two-Stage Joint Assignment Rule",
    59: (
        'The assignment problem introduced in Section 4.1 involves assigning eligible candidate sensors to interceptor '
        'slots of multiple involved targets under capacity constraints. The proposed assignment method enforces that '
        'each sensor can be selected at most once within the current assignment event, while each involved target can '
        'receive no more than its allowed number of interceptors. The method uses a two-stage assignment rule: it '
        'first maximizes the number of filled interceptor slots, and then minimizes the total assignment cost among '
        'assignments with the same number of filled slots.'
    ),
    64: (
        'Table 2 compares the assignments produced by three different selection rules. The greedy assignment processes '
        'the targets sequentially and is therefore order dependent. The MinMax assignment minimizes the worst '
        'target-level team cost, which provides a fairness-oriented solution. The proposed two-stage assignment rule '
        'first fills the maximum feasible number of interceptor slots and then selects the feasible assignment with '
        'the lowest total cost.'
    ),
    65: "Table 2: Assignment comparison under greedy, MinMax, and two-stage assignment methods.",
    68: (
        'The proposed two-stage assignment rule selects  and  for , and  and  for . This assignment fills all four '
        'interceptor slots and achieves the lowest total cost among the feasible assignments shown in Table 2. '
        'Therefore, within the involved targets and eligible candidate sensors of the current assignment event, the '
        'proposed method obtains the total-cost optimal assignment after satisfying the maximum-slot objective.'
    ),
    73: (
        'The MinMax method and the proposed two-stage assignment rule differ fundamentally in their optimization '
        'objectives. The MinMax method minimizes the worst-case target-level team cost, leading to a fairness-oriented '
        'assignment that prevents any involved target from receiving a significantly inferior interceptor configuration. '
        'This objective is meaningful when balanced tracking support across targets is the primary system priority.'
    ),
    74: (
        'In contrast, the proposed two-stage assignment rule first maximizes the number of filled interceptor slots '
        'across the involved targets and then minimizes the total assignment cost among assignments with the same '
        'number of filled slots. This rule prioritizes assignment coverage and total assignment cost minimization '
        'rather than worst-case fairness. Therefore, the two methods may produce different assignments under the same '
        'bidding-cost matrix: MinMax may accept a higher total assignment cost to reduce the worst target-level cost, '
        'while the proposed rule may accept a less balanced distribution of target-level costs if doing so reduces the '
        'total assignment cost after the maximum-slot objective is satisfied. In this report, the two-stage assignment '
        'rule is adopted because the goal is to resolve interceptor assignment conflicts by jointly considering all '
        'targets included in the expanded joint assignment and enforcing non-overlapping sensor use within the current '
        'assignment event.'
    ),
    77: (
        'The proposed two-stage assignment rule preserves the distributed DES architecture by avoiding a permanent '
        'coordinator or centralized bid collector. The method changes the assignment rule, but it does not change the '
        'sensor-level execution logic of the original framework from [1]. Each sensor still operates according to its '
        'own FSM, computes its own bid locally, and updates its own state based on the assignment result. The '
        'optimality of the method refers to the event-local assignment problem over the involved targets and eligible '
        'candidate sensors, rather than centralized control over the entire WSN.'
    ),
    79: (
        'After receiving the request, candidate sensors broadcast their bid values during a fixed bidding window, '
        'denoted by . This window allows bid messages to propagate through the participating sensor group before the '
        'selection step is executed. In the simulation,  can be set to one DES time step, or to a fraction of a time '
        'step, such as half a time step, depending on the desired trade-off between communication completeness and '
        'response speed. At the end of the bidding window, each participating sensor reconstructs the same event-local '
        'assignment problem using the bid information received before the deadline and applies the same deterministic '
        'two-stage assignment rule. The rule is identical for all participating sensors and enforces the same capacity '
        'constraints: each target can receive at most two interceptors, and each sensor can be assigned to at most one '
        'target.'
    ),
    82: (
        'The total delay of the assignment process includes more than the computation time of the two-stage assignment '
        'rule. It also includes the request propagation time, the bidding window, local assignment computation and '
        'commitment broadcast:'
    ),
    88: (
        'The experiments are conducted using the simulation framework developed in [1]. The same discrete-event-system '
        '(DES) architecture, sensor models, EKF-based state estimation, bidding-cost formulation, and target motion '
        'configurations are adopted to ensure a fair comparison between the original MinMax-based method and the '
        'proposed expanded joint assignment approach with a two-stage assignment rule. The main difference between the '
        'two versions lies in how interceptor assignment conflicts are handled once handover requests compete for the '
        'same candidate sensors or when a committed interceptor is contested by a later handover request.'
    ),
    90: (
        'Because the sensor behaviors, target motion, EKF settings, and bidding-cost computation are kept the same in '
        'all experiments, any differences in the assignment outcomes arise from the assignment strategy itself: the '
        'original MinMax-based conflict-resolution method versus the proposed expanded joint assignment approach that '
        'enlarges the joint assignment scope when needed and solves the resulting assignment problem using a two-stage '
        'assignment rule. Sections 5.2 and 5.3 use a two-target scenario to demonstrate the committed-interceptor '
        'conflict and the effect of the expanded joint assignment scope, while Section 5.4 uses a three-target example '
        'to compare the assignment-cost outcomes under MinMax and the proposed two-stage assignment rule.'
    ),
    104: (
        'This case shows that the baseline method does not include the previously served target in the assignment '
        'problem when one of its committed interceptors is contested by a later handover request. Although Target 1 '
        'and Target 2 were not handled within the same assignment procedure, the later handover request from Target 2 '
        'still affected the existing interceptor assignment of Target 1. This motivates the expanded joint assignment '
        'scope introduced in Section 4.1.'
    ),
    106: "Effect of the Expanded Joint Assignment Scope",
    107: (
        'The proposed expanded joint assignment approach addresses the committed-interceptor conflict by enlarging the '
        'joint assignment scope before solving the assignment problem. When the later handover request from Target 2 '
        'attempts to use , the system detects that  has already been committed to Target 1. Therefore, Target 1 is '
        'added to the expanded joint assignment together with the current calling target, Target 2.'
    ),
    108: (
        'The assignment problem is then solved once over the expanded joint assignment and the eligible candidate '
        'sensors. This prevents the later handover request from directly overwriting the existing interceptor assignment '
        'of Target 1. Instead, both targets are considered within the same assignment step, and the two-stage '
        'assignment rule selects a feasible non-overlapping sensor allocation.'
    ),
    110: "Figure 4: Proposed expanded joint assignment scope at s.",
    111: (
        'Figure 4 shows the result after applying the proposed expanded joint assignment approach. Target 2 no longer '
        'takes  away from Target 1. Instead, Target 2 selects  and , while Target 1 retains  as one of its committed '
        'interceptors. This confirms that the proposed method includes the previously served target in the assignment '
        'problem when its committed interceptor is contested by a later handover request.'
    ),
    113: "Effect of the Expanded Joint Assignment Scope and Assignment Rule in a Three-Target Event",
    114: (
        'This section uses a three-target event to separate two effects: the expanded joint assignment scope and '
        'assignment-rule selection. The original MinMax-based method without the expanded joint assignment scope is '
        'discussed as a scope limitation, because it would not form the full [, , ] assignment problem as a single '
        'event-level instance. Therefore, the event-level cost comparison is made only between methods that use the '
        'same expanded joint assignment scope: MinMax with expanded joint assignment scope and the proposed two-stage '
        'assignment rule with expanded joint assignment scope.'
    ),
    115: (
        'Figure 5 illustrates the three-target expanded joint assignment event. Target 3 first issues a handover '
        'request. During the assignment process,   is found to have already been committed to Target 2, so Target 2 is '
        'added to the expanded joint assignment. A further committed-interceptor overlap occurs when the assignment '
        'process involving Target 2 includes , which has already been committed to Target 1. As a result, the '
        'assignment scope expands from a single calling target to the three-target joint assignment [, , ]. This '
        'example shows that the expanded joint assignment scope prevents previously served targets from being excluded '
        'when their committed interceptors are contested by later handover requests.'
    ),
    117: "Figure 5: Three-target expanded joint assignment event at s.",
    118: (
        'Once the expanded joint assignment scope is fixed, MinMax and the proposed two-stage assignment rule differ '
        'only in their assignment objective. MinMax prioritizes the worst target-level team cost, while the proposed '
        'two-stage assignment rule first maximizes the number of filled interceptor slots and then minimizes the total '
        'assignment cost.'
    ),
    121: (
        'Table 3 shows that both methods fill all six requested interceptor slots, but the two-stage assignment rule '
        'achieves a lower total assignment cost, reducing the cost from 1.47187 to 1.33343. This result supports the '
        'event-level rule defined in this report: after the maximum-slot objective is satisfied, the two-stage '
        'assignment rule selects the feasible non-overlapping assignment with the lowest total cost.'
    ),
    124: (
        'Table 4 provides a complementary observation rather than the primary event-level comparison. In this '
        'simulation seed, the two-stage assignment rule produces a slightly lower total tracked-time ratio than the '
        'MinMax variant, 31.94% compared with 32.59%. The per-target results also show that the two-stage rule shifts '
        'more tracking time toward Target 1, while Target 2 and Target 3 receive less tracked time than under MinMax. '
        'This does not contradict the event-level assignment result, because full-simulation tracking time depends on '
        'later target motion, subsequent handover timing, future conflicts, and downstream effects of earlier '
        'assignment decisions.'
    ),
    125: (
        'Overall, this example shows that the expanded joint assignment scope and assignment-rule selection should be '
        'understood separately. The expanded joint assignment scope determines which targets are included in the '
        'assignment problem, while the choice between MinMax and the two-stage assignment rule determines how the '
        'system prioritizes balanced target-level support versus lower event-level total assignment cost.'
    ),
    128: (
        'This report studied interceptor assignment conflicts in a mobile WSN for cooperative multi-target tracking. '
        'Building on the DES-based proactive handover framework in [1], the report focused on the case where a later '
        'handover request contests a sensor that has already been committed as an interceptor for another target. To '
        'address this issue, the proposed expanded joint assignment approach includes not only the current calling '
        'targets but also any previously served target whose committed interceptor is contested in the same assignment '
        'problem. The resulting assignment is solved using a two-stage assignment rule that first maximizes the number '
        'of filled interceptor slots and then minimizes the total assignment cost. Simulation examples show that the '
        'expanded joint assignment scope prevents previously served targets from being excluded during committed-'
        'interceptor conflicts, while the two-stage assignment rule achieves lower event-level total assignment cost '
        'than MinMax in the examined three-target event. The results also show that MinMax remains a meaningful '
        'fairness-oriented baseline, while the proposed expanded joint assignment approach prioritizes event-level '
        'total assignment cost minimization.'
    ),
    130: (
        'Future work should evaluate the proposed expanded joint assignment approach under more realistic communication '
        'and hardware constraints. The current simulation assumes ideal communication, but physical deployment would '
        'require measuring bid propagation delay, packet loss, commitment confirmation time, and the possibility of '
        'inconsistent assignment reconstruction among sensors. In addition, the simulation results suggest that lower '
        'event-level assignment cost does not necessarily guarantee better aggregate tracked-target time in every seed. '
        'Future work should therefore investigate adaptive rule selection, where the system can switch between MinMax '
        'and the two-stage assignment rule depending on target priority, uncertainty level, fairness requirements, or '
        'long-horizon tracking performance. More simulation seeds and more diverse multi-target trajectories should '
        'also be tested to evaluate the robustness of the expanded joint assignment scope under different conflict '
        'patterns.'
    ),
    135: (
        'The appendix includes the main simulation script and representative log files used for the expanded joint '
        'assignment experiments.'
    ),
    136: "Example_V47_Yeqi.m бк main simulation script for the V47 expanded joint assignment experiment.",
    137: (
        'example2-affected-group-lexicographic.log бк log file for the expanded joint assignment experiment using the '
        'two-stage assignment rule.'
    ),
    138: "example2-affected-group-minmax.log бк log file for the expanded joint assignment experiment using the MinMax assignment objective.",
}


TABLE_UPDATES = {
    (1, 3, 0): "Two-stage rule",
    (2, 1, 0): "MinMax with expanded joint assignment scope",
    (2, 2, 0): "Two-stage rule with expanded joint assignment scope",
    (3, 1, 0): "MinMax with expanded joint assignment scope",
    (3, 2, 0): "Two-stage rule with expanded joint assignment scope",
}


TEXT_REPLACEMENTS = [
    (
        "This extension ensures that the previously served target is also included in the assignment problem",
        "This extension ensures that the previously served target is also included in the expanded joint assignment",
    ),
    (
        "Following the affected-target grouping mechanism in Section 4.1",
        "Following the expanded joint assignment scope defined in Section 4.1",
    ),
    (
        "the assignment must be solved jointly over the affected-target set while enforcing non-overlapping sensor use.",
        "the assignment must be solved jointly over the involved targets while enforcing non-overlapping sensor use.",
    ),
]


TOKEN_RE = re.compile(r"\w+|\s+|[^\w\s]", re.UNICODE)


@dataclass
class ParagraphBlock:
    text: str
    style: str


@dataclass
class TableBlock:
    rows: list[list[str]]


def iter_block_items(parent: _Document):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_blocks(doc: _Document):
    blocks = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            blocks.append(ParagraphBlock(text=block.text, style=block.style.name))
        else:
            rows = [[cell.text for cell in row.cells] for row in block.rows]
            blocks.append(TableBlock(rows=rows))
    return blocks


def update_paragraph_text(doc: _Document, index: int, new_text: str) -> None:
    doc.paragraphs[index].text = new_text


def update_table_cell(doc: _Document, table_idx: int, row_idx: int, col_idx: int, new_text: str) -> None:
    doc.tables[table_idx].cell(row_idx, col_idx).text = new_text


def apply_text_replacements(doc: _Document) -> None:
    for para in doc.paragraphs:
        text = para.text
        new_text = text
        for old, new in TEXT_REPLACEMENTS:
            new_text = new_text.replace(old, new)
        if new_text != text:
            para.text = new_text


def diff_html(old: str, new: str) -> str:
    if old == new:
        return escape(new)
    old_tokens = TOKEN_RE.findall(old)
    new_tokens = TOKEN_RE.findall(new)
    matcher = SequenceMatcher(a=old_tokens, b=new_tokens)
    parts = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        old_chunk = escape("".join(old_tokens[i1:i2]))
        new_chunk = escape("".join(new_tokens[j1:j2]))
        if tag == "equal":
            parts.append(new_chunk)
        elif tag == "insert":
            parts.append(f"<mark>{new_chunk}</mark>")
        elif tag == "delete":
            parts.append(f"<del>{old_chunk}</del>")
        elif tag == "replace":
            if old_chunk:
                parts.append(f"<del>{old_chunk}</del>")
            if new_chunk:
                parts.append(f"<mark>{new_chunk}</mark>")
    return "".join(parts)


def paragraph_tag(block_index: int, block: ParagraphBlock) -> str:
    text = block.text.strip()
    if block_index == 0:
        return "h1"
    if block_index in (1, 2):
        return "p class='meta'"
    if text == "Abstract":
        return "h2"
    if block.style == "List Paragraph":
        return "h2"
    return "p"


def render_html(original_blocks, updated_blocks) -> str:
    assert len(original_blocks) == len(updated_blocks)
    body_parts = []
    for idx, (old_block, new_block) in enumerate(zip(original_blocks, updated_blocks)):
        if isinstance(old_block, ParagraphBlock) and isinstance(new_block, ParagraphBlock):
            tag = paragraph_tag(idx, new_block)
            body_parts.append(f"<{tag}>{diff_html(old_block.text, new_block.text)}</{tag.split()[0]}>")
        elif isinstance(old_block, TableBlock) and isinstance(new_block, TableBlock):
            rows_html = []
            for old_row, new_row in zip(old_block.rows, new_block.rows):
                cells_html = []
                for old_cell, new_cell in zip(old_row, new_row):
                    cells_html.append(f"<td>{diff_html(old_cell, new_cell)}</td>")
                rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
            body_parts.append("<table>" + "".join(rows_html) + "</table>")
    style = """
    body { font-family: Georgia, 'Times New Roman', serif; margin: 40px auto; max-width: 920px; line-height: 1.6; color: #222; }
    h1 { font-size: 30px; margin-bottom: 8px; }
    h2 { font-size: 20px; margin-top: 28px; margin-bottom: 8px; }
    p { margin: 10px 0; }
    .meta { margin: 2px 0; color: #555; }
    table { border-collapse: collapse; width: 100%; margin: 18px 0; font-size: 14px; }
    td, th { border: 1px solid #bbb; padding: 8px 10px; vertical-align: top; }
    mark { background: #fff2a8; padding: 0 1px; }
    del { background: #ffe0e0; color: #8a1f1f; text-decoration: line-through; }
    .legend { margin: 20px 0 28px; padding: 12px 14px; background: #f7f7f7; border-left: 4px solid #999; font-size: 14px; }
    """
    legend = (
        "<div class='legend'><strong>Highlight guide:</strong> "
        "<mark>yellow</mark> shows new wording, and <del>red strikethrough</del> shows removed wording.</div>"
    )
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<title>Expanded Joint Assignment Report - Highlighted Changes</title>"
        f"<style>{style}</style></head><body>{legend}{''.join(body_parts)}</body></html>"
    )


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    original_doc = Document(SOURCE_DOCX)
    original_blocks = extract_blocks(original_doc)

    updated_doc = Document(SOURCE_DOCX)
    for index, new_text in PARA_UPDATES.items():
        update_paragraph_text(updated_doc, index, new_text)
    for table_idx, row_idx, col_idx, new_text in (
        (table_idx, row_idx, col_idx, new_text)
        for (table_idx, row_idx, col_idx), new_text in TABLE_UPDATES.items()
    ):
        update_table_cell(updated_doc, table_idx, row_idx, col_idx, new_text)
    apply_text_replacements(updated_doc)
    updated_doc.save(OUTPUT_DOCX)

    final_doc = Document(OUTPUT_DOCX)
    updated_blocks = extract_blocks(final_doc)
    html = render_html(original_blocks, updated_blocks)
    OUTPUT_HTML.write_text(html, encoding="utf-8")

    print(OUTPUT_DOCX)
    print(OUTPUT_HTML)


if __name__ == "__main__":
    main()
