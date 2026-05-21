from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path(r"C:\Users\abel\Desktop\resume\Yeqi_Sang_2025_cv.docx")
OUT_DIR = Path(r"C:\Users\abel\Desktop\MEng\MIE8888\mie8888\output\doc")
OUT_PATH = OUT_DIR / "Yeqi_Sang_2025_cv_updated_draft.docx"
PREVIEW_PATH = OUT_DIR / "Yeqi_Sang_2025_cv_updated_draft_preview.txt"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def set_paragraph_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def find_paragraph(doc, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    raise ValueError("Paragraph not found")


def paragraph_index(doc, target):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph._element is target._element:
            return idx
    raise ValueError("Paragraph index not found")


def replace_paragraph_block(doc, start_paragraph, end_paragraph, new_texts):
    start_idx = paragraph_index(doc, start_paragraph)
    end_idx = paragraph_index(doc, end_paragraph)
    block = doc.paragraphs[start_idx:end_idx]
    if len(block) < len(new_texts):
        raise ValueError("Replacement block is shorter than requested text list")

    for paragraph, text in zip(block, new_texts):
        set_paragraph_text(paragraph, text)

    for paragraph in list(block[len(new_texts):]):
        remove_paragraph(paragraph)


def append_preview(doc):
    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    PREVIEW_PATH.write_text("\n".join(lines), encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT_PATH)
    doc = Document(OUT_PATH)

    contact = find_paragraph(doc, lambda p: p.text.strip().startswith("Email:"))
    set_paragraph_text(
        contact,
        "Email: abel.sang@mail.utoronto.ca | Phone: (+1) 647-937-7099 | Toronto, ON",
    )

    objective_heading = find_paragraph(doc, lambda p: p.text.strip() == "Objective")
    set_paragraph_text(objective_heading, "Professional Summary")

    summary = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Results-driven DevOps Engineer / Software Engineer"),
    )
    set_paragraph_text(
        summary,
        "Data Scientist with 5 years of experience across software engineering, cloud infrastructure, and applied analytics. Currently supporting AML client risk model benchmarking and model performance improvement initiatives, with graduate project experience in state estimation, optimization, and simulation-based validation.",
    )

    highlights_heading = find_paragraph(
        doc, lambda p: p.text.strip() == "Highlights of Qualifications"
    )
    set_paragraph_text(highlights_heading, "Core Qualifications")

    first_highlight = find_paragraph(
        doc,
        lambda p: p.text.strip() == "5 years of experience in DevOps, full-stack development, and cloud computing.",
    )
    work_experience_heading = find_paragraph(doc, lambda p: p.text.strip() == "Work Experience")
    replace_paragraph_block(
        doc,
        first_highlight,
        work_experience_heading,
        [
            "5 years of experience spanning data science, software engineering, and cloud-based systems.",
            "Hands-on experience with model benchmarking, performance analysis, and structured problem solving in analytics-driven environments.",
            "Strong technical foundation in Python, SQL, MATLAB, Node.js, and cloud tooling including AWS, Docker, and CI/CD pipelines.",
            "Graduate project experience in state estimation, optimization, simulation, and algorithm validation.",
            "Effective cross-functional collaborator with experience partnering with stakeholders and supporting shared VM and server environments.",
        ],
    )

    work_experience_heading = find_paragraph(doc, lambda p: p.text.strip() == "Work Experience")
    scotia_title = insert_paragraph_after(
        work_experience_heading,
        "Data Scientist, Scotiabank, Toronto, Canada                                                                                                   [Start Date] - Present",
        style="Normal",
    )
    scotia_bullet_1 = insert_paragraph_after(
        scotia_title,
        "Built benchmarking analyses for AML client risk models to evaluate current-model performance, compare approaches, and identify opportunities for improvement.",
        style="List Paragraph",
    )
    scotia_bullet_2 = insert_paragraph_after(
        scotia_bullet_1,
        "Partnered with business and technical stakeholders to gather requirements, communicate findings, and support model enhancement initiatives in the Anti-Money Laundering domain.",
        style="List Paragraph",
    )
    insert_paragraph_after(
        scotia_bullet_2,
        "Managed team virtual machines and server environments to maintain reliable infrastructure for analytics and development workflows.",
        style="List Paragraph",
    )

    allied_b1 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Integrated Snowflake LLM and AI pipelines"),
    )
    replace_paragraph_block(
        doc,
        allied_b1,
        find_paragraph(doc, lambda p: p.text.strip().startswith("Software Engineer, Orbiseed Technology Inc")),
        [
            "Integrated Snowflake-based LLM and AI pipelines that routed applications to different models by user role, improving scalability and data efficiency.",
            "Maintained and modernized legacy systems inherited through acquisition, preserving system stability and operational continuity.",
            "Automated lower-environment hibernation workflows to reduce cloud infrastructure costs while maintaining development availability.",
            "Managed end-to-end CI/CD pipelines, virtual machines, and dynamic proxy provisioning to improve deployment reliability and support operational needs.",
        ],
    )

    orbiseed_b1 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Cooperated with Algorithm Engineers in-depth"),
    )
    orbiseed_heading = find_paragraph(doc, lambda p: p.text.strip() == "Education")
    replace_paragraph_block(
        doc,
        orbiseed_b1,
        orbiseed_heading,
        [
            "Collaborated with algorithm engineers to develop and productionize a Python and OpenCV pipeline for generating 3D building models from 2D inputs with fewer corner cases.",
            "Built a web-based 3D viewer demo with Three.js and WebGL in one month to showcase Building Information System outputs and support investor-facing presentations.",
            "Led Lunch-and-Learn sessions to share new technologies and practical engineering learnings across teams.",
        ],
    )

    research_heading = find_paragraph(doc, lambda p: p.text.strip() == "Research")
    set_paragraph_text(research_heading, "Projects & Research")
    project_title = insert_paragraph_after(
        research_heading,
        "Master of Engineering Project - Cooperative Multi-Target Tracking in Mobile Wireless Sensor Networks, University of Toronto                     2025",
        style="Normal",
    )
    project_b1 = insert_paragraph_after(
        project_title,
        "Developed a MATLAB simulation for cooperative multi-target tracking in a mobile wireless sensor network, coordinating a 25-node sensor swarm to track up to 3 moving targets.",
        style="List Paragraph",
    )
    project_b2 = insert_paragraph_after(
        project_b1,
        "Implemented per-sensor Extended Kalman Filter (EKF) state estimation and proactive handover logic in a distributed swarm system, enabling the network to predict target loss and dispatch interceptors before tracking failure.",
        style="List Paragraph",
    )
    insert_paragraph_after(
        project_b2,
        "Designed and extended conflict-resolution algorithms for simultaneous handover events in a distributed event-driven system, introducing global assignment logic and validating improvements through simulation-based A/B testing against legacy strategies.",
        style="List Paragraph",
    )

    ra_b1 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Contributed to a CHI 2025-accepted paper"),
    )
    ra_b2 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Developed and implemented a VR environment for system testing and evaluation."),
    )
    ra_b3 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Designed and integrated a coordinate system mapping framework"),
    )
    set_paragraph_text(
        ra_b1,
        'Contributed to the CHI 2025-accepted paper "VibraForge: A Scalable Prototyping Toolkit for Creating Spatialized Vibrotactile Feedback Systems."',
    )
    set_paragraph_text(
        ra_b2,
        "Developed a VR testing environment and a coordinate-mapping framework to align virtual and real-world tracking systems for evaluation.",
    )
    remove_paragraph(ra_b3)

    doc.save(OUT_PATH)
    append_preview(doc)
    print(f"Saved draft to: {OUT_PATH}")
    print(f"Saved preview to: {PREVIEW_PATH}")


if __name__ == "__main__":
    main()
