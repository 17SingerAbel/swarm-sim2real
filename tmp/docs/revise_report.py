from __future__ import annotations

import html
import re
import shutil
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document


ROOT = Path(r"C:/Users/abel/Desktop/MIE8888/mie8888")
SOURCE_DOC = Path(r"C:/Users/abel/Desktop/YS2_NN1_mie8888-final-report.docx")
OUTPUT_DOC = ROOT / "output/doc/YS2_NN1_mie8888-final-report_revised_global_assignment.docx"
OUTPUT_HTML = ROOT / "output/doc/YS2_NN1_mie8888-final-report_changes.html"


@dataclass(frozen=True)
class ParagraphEdit:
    index: int
    label: str
    expected: str
    new_text: str


@dataclass(frozen=True)
class TableEdit:
    table_idx: int
    row_idx: int
    col_idx: int
    label: str
    expected: str
    new_text: str


PARAGRAPH_EDITS = [
    ParagraphEdit(
        5,
        "Abstract paragraph 1",
        "This report studies interceptor assignment",
        "This report studies interceptor assignment in a mobile wireless sensor network (WSN) used for cooperative multi-target tracking. The underlying framework already uses a discrete-event-system (DES) / finite-state-machine (FSM) architecture with handover request broadcasting, local bid-cost computation by eligible candidate sensors, and conflict resolution during interceptor selection. In the original simulation (V45), conflict resolution was based on a MinMax objective that attempts to minimize the worst team cost among the requested targets. This fairness-oriented objective protects the weakest target assignment, but it does not directly minimize the total assignment cost when different targets compete for the same eligible sensors.",
    ),
    ParagraphEdit(
        6,
        "Abstract paragraph 2",
        "The proposed (V46) assignment reformulates the problem as a global assignment problem",
        "The proposed assignment method in V46 reformulates this conflict-resolution stage as a joint global assignment over all affected targets and all eligible candidate sensors considered in the current assignment process. It also introduces a joint reassignment mechanism for the DES case in which a later request competes with interceptors already committed to another target. The proposed assignment first maximizes the total number of assigned interceptors across the affected targets and then, among assignments with the same interceptor coverage, minimizes the total bid cost.",
    ),
    ParagraphEdit(
        7,
        "Abstract paragraph 3",
        "The report compares the proposed method and the minimax method",
        "The report compares the proposed assignment method and the original MinMax method using the same affected targets, eligible candidate sensors, and bid costs under the same simulation communication assumptions. A natural simulation conflict is used as the primary worked example. In that event, the two methods produce different feasible assignments because they optimize different objectives: MinMax is fairness-oriented, whereas the proposed assignment method first seeks to assign as many interceptors as possible across the affected targets and then minimizes the total bidding cost. The results therefore support an algorithmic comparison between the two formulations and show how joint reassignment can prevent conflict-induced loss of previously committed interceptors, while the communication and deployment implications of global coordination are discussed explicitly as limitations for future hardware validation.",
    ),
    ParagraphEdit(
        13,
        "Introduction contribution paragraph",
        "The original conflict-resolution method used a minimax objective.",
        "The original conflict-resolution method used a minimax objective. When affected targets compete for the same candidate sensors, the minimax method attempted to select teams such that the worst target team cost was minimized. This is a fairness-oriented objective: no target should receive a very poor team if a more balanced solution is available. However, conflict is not limited to requests handled within the same assignment procedure. A conflict also occurs when a later request competes with interceptors already committed to another target. In V46, the proposed assignment method reformulates interceptor selection as a joint global assignment over all affected targets and all eligible candidate sensors, and introduces a joint reassignment mechanism for this broader DES conflict condition. The proposed assignment first maximizes the number of fulfilled interceptor requests across affected targets and then minimizes the total assignment cost, subject to the constraints that each target receives at most two interceptors and each sensor is assigned to at most one target.",
    ),
    ParagraphEdit(
        14,
        "Introduction bullet 1",
        "It reformulates interceptor selection as a global assignment problem",
        "It reformulates interceptor selection as a joint global assignment over all affected targets and all eligible candidate sensors.",
    ),
    ParagraphEdit(
        15,
        "Introduction bullet 2",
        "It compares global assignment framework and MinMax assignment",
        "It introduces a joint reassignment mechanism when later requests compete with interceptors already committed to another target.",
    ),
    ParagraphEdit(
        16,
        "Introduction bullet 3",
        "It uses natural simulation evidence to show how the two objectives can make different decisions.",
        "It compares the proposed assignment method and MinMax using the same candidate sensors and bid costs, showing how the two objectives can make different decisions.",
    ),
    ParagraphEdit(
        17,
        "Introduction bullet 4",
        "It analyzes feasibility and communication limitations",
        "It analyzes feasibility, communication assumptions, and latency limitations for Crazyflie 2.0 deployment.",
    ),
    ParagraphEdit(
        24,
        "Related Work closing paragraph",
        "This report builds upon that framework and focuses specifically",
        "This report builds upon that framework and focuses specifically on interceptor assignment conflicts. The proposed assignment method in V46 introduces a joint global assignment over all affected targets and all eligible candidate sensors whenever they compete for shared interceptor resources, together with a joint reassignment mechanism when later requests conflict with already committed interceptors. Unlike the original min-max method, which balances the worst-case team cost, the proposed method first maximizes feasible interceptor allocation and then minimizes total assignment cost, providing an alternative objective for system-level resource allocation.",
    ),
    ParagraphEdit(
        30,
        "Section 3 system assumptions",
        "The system assumptions used in the remainder of the report are as follows",
        "The system assumptions used in the remainder of the report are as follows: the workspace is planar; sensors are homogeneous; assignment is event-triggered when a target enters the interceptor assignment process; and active trackers are excluded from candidate bidding during the current event. Each interceptor assignment event follows three steps: broadcast, bidding and selection, corresponding to the states PENDING_BROADCAST, PENDING_BIDDING and PENDING_SELECTION. Since different targets may enter this process at different times, a new request may compete with interceptors already committed to another target. If the requests share one or more candidate sensors, the system treats this as one interceptor assignment conflict and resolves it through a joint global assignment over the affected targets. In this report, global assignment refers to a joint assignment over all eligible candidate sensors and all affected targets considered in the current assignment process. It does not imply continuous centralized control over the entire WSN at every time step.",
    ),
    ParagraphEdit(
        60,
        "Section 4.2 opening paragraph",
        "The joint assignment problem introduced in Section 4.1",
        "The joint assignment problem introduced in Section 4.1 involves assigning all eligible candidate sensors to interceptor slots of all affected targets under capacity constraints. In this report, the proposed assignment method in V46 jointly considers these sensors and targets using the bid costs defined in Section 3.2. It first maximizes the total number of assigned interceptors across the affected targets and then, among assignments with the same number of assigned interceptors, minimizes the total assignment cost. Each sensor can be used at most once, and each target can receive at most two interceptors.",
    ),
    ParagraphEdit(
        64,
        "Section 4.2 example explanation",
        "Each candidate sensor can be selected at most once.",
        "Each candidate sensor can be selected at most once. Therefore, the algorithm cannot simply choose the two lowest-cost sensors for each target independently, because the same sensor may be selected by more than one target based on the bidding result. Instead, the proposed assignment method performs a joint global assignment over all affected targets and all eligible candidate sensors in the current assignment process, using the available bid costs and respecting the per-sensor and per-target capacity limits. Figure 2 illustrates the sensor-target assignment structure for interceptor selection.",
    ),
    ParagraphEdit(
        65,
        "Table 2 caption paragraph",
        "Table 2. Result Assignment Table from Greedy, MinMax and Global Methods",
        "Table 2. Result Assignment Table from Greedy, MinMax, and Global Assignment Methods",
    ),
    ParagraphEdit(
        66,
        "Table 2 explanation paragraph",
        "Table 2 provides a comparison for assignment results",
        "Table 2 provides a comparison for assignment results from the Greedy method, MinMax, and the proposed assignment method. The greedy method produces its assignment because the targets are processed sequentially in the order [2, 1], which is induced by the reassignment mechanism triggered after the conflict is detected. As a result, Target 2 is handled first, and Target 1 is processed afterward. The greedy result is therefore order-dependent and based on local decisions rather than joint optimization. By contrast, the MinMax method jointly considers both targets and aims to minimize the maximum team cost among them. This means that even if assigning S10 and S14 to Target 1 yields a lower team cost than assigning S14 and S18, both assignments are equally optimal under an exact MinMax formulation if they lead to the same worst-case team cost. Finally, the proposed assignment method considers both targets and all eligible candidate sensors simultaneously within the joint reassignment event. Because the compared assignments all satisfy the same interceptor coverage in this example, the second-level objective determines the result and selects the feasible assignment with the minimum total cost. In this case, it assigns S13 and S20 to Target 2 and S10 and S14 to Target 1, thereby achieving the lowest total cost among the candidate joint assignments.",
    ),
    ParagraphEdit(
        68,
        "Section 4.3 heading",
        "Objective Trade-off: MinMax v.s. Global Assignment",
        "Objective Trade-off: MinMax vs. Global Assignment",
    ),
    ParagraphEdit(
        69,
        "Section 4.3 paragraph 1",
        "The MinMax and global assignment method approaches differ fundamentally",
        "The MinMax and proposed global assignment approaches differ fundamentally in their optimization objectives. The MinMax method minimizes the worst-case target cost, leading to a fairness-oriented assignment that prevents any target from receiving a significantly inferior interceptor configuration. This objective is particularly beneficial when balanced tracking performance across all targets is critical.",
    ),
    ParagraphEdit(
        70,
        "Section 4.3 paragraph 2",
        "In contrast, the global assignment formulation minimizes",
        "In contrast, the proposed global assignment formulation assigns individual sensors to targets jointly across all affected targets. Its first priority is to maximize the number of assigned interceptors under feasibility and capacity constraints; only among assignments with the same coverage does it minimize the total assignment cost. In other words, it seeks to provide as much interceptor support as possible across the affected targets and then minimize the total bidding cost, but it does not explicitly enforce fairness among targets. Consequently, the two methods may produce different assignments under the same cost matrix: MinMax may sacrifice lower total cost to improve the worst-case target, while the proposed global assignment may accept uneven target costs if doing so reduces the total system cost after feasible coverage has been maximized.",
    ),
    ParagraphEdit(
        71,
        "Section 4.3 paragraph 3",
        "The two objectives reflect different system priorities.",
        "The two objectives reflect different system priorities. MinMax remains a meaningful fairness-oriented baseline because it explicitly protects the worst-case target. The proposed assignment method is adopted here as an alternative rule that first maximizes feasible interceptor allocation and then minimizes total bidding cost, not as a universally better method. Under the same bid costs, the two objectives can legitimately produce different assignments.",
    ),
    ParagraphEdit(
        75,
        "Section 4.4.1 transition sentence",
        ", t,",
        "This section focuses on what is assumed at the assignment step in simulation and what would still need to be resolved for a practical distributed deployment.",
    ),
    ParagraphEdit(
        76,
        "Section 4.4.1 paragraph 1",
        "The proposed global assignment method preserves the distributed DES architecture",
        "The proposed assignment method preserves the distributed DES architecture. At the system level, handover requests are broadcast, eligible candidate sensors compute local bid costs, and the assignment decision is formed from the available bids. In V46, when a conflict involves sensors already committed to another target, the affected targets are grouped into one joint reassignment process. The proposed assignment method then solves a joint global assignment over all affected targets and all eligible candidate sensors, subject to the constraints that each target receives at most two interceptors and each sensor is assigned to at most one target.",
    ),
    ParagraphEdit(
        77,
        "Section 4.4.1 paragraph 2",
        "In a practical distributed implementation",
        "In this report, global assignment refers to a joint assignment over all eligible candidate sensors and all affected targets considered in the current assignment process. It does not imply continuous centralized control over the entire WSN at every time step. Under the same global-broadcast / perfect-communication assumption used in the original simulation, bid costs from all eligible candidates are assumed to be available at the selection step. A real distributed hardware implementation would still require either an explicit final-assignment broadcast or a deterministic consensus or reconstruction mechanism so that all participating sensors agree on the same final assignment.",
    ),
    ParagraphEdit(
        84,
        "Section 4.4.2 closing paragraph",
        "The cost table is also small",
        "The cost table is also small, with at most M x N bidding costs. Therefore, the assignment problem is bounded for the current simulation. For larger systems, optional deployment-oriented shortcuts such as limiting the candidate sensor set could reduce computation or communication, but they are not the main method studied here. The main formulation assumes a joint global assignment over all eligible candidate sensors in the current assignment process.",
    ),
    ParagraphEdit(
        87,
        "Section 4.4.3 computation paragraph",
        "Using the workstation measurement as an optimistic computation-only baseline",
        "Using the workstation measurement as an optimistic computation-only baseline, the global assignment computation time of 0.001254 s on a 4.5 GHz CPU corresponds to about 5.64 x 10^6 processor cycles. Under ideal frequency-only scaling, the same cycle count would require about 0.0336 s on the 168 MHz STM32F405, or 33.6% of a 0.1 s simulation time step. However, this estimate addresses assignment computation only and is not sufficient to establish real-time feasibility, because embedded execution also depends on communication latency, synchronization delay, packet loss handling, final assignment confirmation, memory effects, floating-point efficiency, and concurrent control workload.",
    ),
    ParagraphEdit(
        89,
        "Section 4.4.3 latency paragraph",
        "More importantly, the sensor-to-sensor response budget must include both computation and communication.",
        "More importantly, the total response budget for one assignment event must include the full communication-and-decision chain, not only the assignment computation itself. A more complete latency model is T_total = T_request_broadcast + T_bid_computation + T_bid_transmission + T_assignment + T_result_broadcast + T_commit_update. Under the same global-broadcast / perfect-communication assumption used in the current simulation, the global assignment result remains meaningful for algorithmic comparison. However, real embedded feasibility would require communication-aware validation of the entire event pipeline, including request broadcast from the active tracker, local bid computation by all eligible candidate sensors, bid transmission or bid collection, assignment computation, final assignment broadcast, and commitment or state update by selected and non-selected sensors. Therefore, computation time alone is not sufficient to prove real-time feasibility.",
    ),
    ParagraphEdit(
        93,
        "Section 5.1 experimental setup paragraph 1",
        "The experiments are conducted using the simulation framework",
        "The experiments are conducted using the simulation framework developed in [1]. The same DES/FSM architecture, sensor models, target motion configurations, request broadcasting, and local bid-cost formulation are retained to ensure a fair comparison between the original MinMax-based assignment (V45) and the proposed assignment method (V46).",
    ),
    ParagraphEdit(
        95,
        "Section 5.1 experimental setup paragraph 3",
        "The same EKF-based state estimation and bidding cost formulation",
        "The same EKF-based state estimation and bidding cost formulation are used in both methods. This ensures that any differences in assignment results arise solely from the assignment objective and the joint reassignment logic (MinMax versus the proposed assignment method), rather than differences in the underlying sensing, prediction, or bidding assumptions.",
    ),
    ParagraphEdit(
        106,
        "Section 5.2 conflict-resolution paragraph",
        "To address this issue, a new conflict-management mechanism",
        "To address this issue, the proposed assignment method in V46 introduces a joint reassignment mechanism. When a target attempts to select an interceptor that is already assigned to another target, the system treats the affected targets as one interceptor assignment conflict. Instead of allowing the new target to simply take the sensor away, the affected targets are grouped together, and one joint global assignment is performed over all affected targets and all eligible candidate sensors. If a sensor is reassigned, the original target is simultaneously reconsidered within the same assignment process, so alternative interceptors can be allocated consistently under the same constraints. This allows the final interceptor allocation to be determined jointly for the affected target group rather than target by target.",
    ),
    ParagraphEdit(
        110,
        "Figure 7 caption",
        "Figure7: t=40.1 for V46 simulation.",
        "Figure7: t=40.1 for V46 simulation. Target 2 now employs Sensor 13 and Sensor 20 instead of taking over Sensor 14 from Target 1",
    ),
    ParagraphEdit(
        111,
        "Section 5.2 V46 result paragraph",
        "Figure 7 shows the result after applying the V46 fix.",
        "Figure 7 shows the result after applying the proposed assignment method in V46. At t = 39.5 s, Target 2 no longer takes Sensor 14 away from Target 1. Instead, the updated global assignment assigns Sensor 13 and Sensor 20 to Target 2, while Target 1 retains its assigned interceptors. This confirms that the proposed conflict-resolution logic prevents an active interception assignment from being overwritten by a later target selection under the same bid costs considered in the joint assignment step.",
    ),
    ParagraphEdit(
        113,
        "Section 5.3 heading",
        "MinMax V.S. MCMF on Interceptor Assignment",
        "MinMax vs. Global Assignment on Interceptor Assignment",
    ),
    ParagraphEdit(
        114,
        "Section 5.3 introduction paragraph",
        "This section compares the assignment results produced by MinMax and MCMF",
        "This section compares the assignment results produced by MinMax and the proposed assignment method under the same conflict scenario. The goal is not to claim that one method is universally better, but to illustrate how different optimization objectives lead to different assignment outcomes.",
    ),
    ParagraphEdit(
        118,
        "Section 5.3 figure 8 paragraph",
        "Figure 8 shows the full simulation result from V46",
        "Figure 8 shows the full simulation result from V46 after introducing the updated conflict-resolution mechanism. At t = 39.5 s, Target 1 and Target 2 trigger an assignment conflict, causing the system to perform a joint reassignment process over the affected targets and eligible candidate sensors.",
    ),
    ParagraphEdit(
        119,
        "Table 4 caption paragraph",
        "Table 4: Interceptor Assignment Comparison Between MinMax and MCMF",
        "Table 4: Interceptor Assignment Comparison Between MinMax and Global Assignment",
    ),
    ParagraphEdit(
        120,
        "Table 4 explanation paragraph",
        "Table 4 summarizes the interceptor assignment results produced by MinMax and MCMF",
        "Table 4 summarizes the interceptor assignment results produced by MinMax and the proposed assignment method for the conflict event at t = 39.5. In this example, the proposed assignment method reduces the total assignment cost from 1.181 to 1.084 while maintaining the same total number of assigned interceptors. The difference arises from the distinct optimization objectives. MinMax prioritizes balancing the worst-case target cost, whereas the proposed global assignment first maximizes feasible interceptor allocation and then minimizes the total assignment cost. As a result, the two methods select different sensor combinations under the same scenario.",
    ),
    ParagraphEdit(
        123,
        "Conclusion paragraph",
        "This report reformulates interceptor assignment in cooperative multi-target tracking as a conflict-aware min-cost max-flow problem.",
        "This report proposes a global assignment and joint reassignment method for interceptor conflict resolution in cooperative multi-target tracking. Under the DES/FSM framework, V46 jointly considers all affected targets and all eligible candidate sensors in the current assignment process. The proposed assignment first maximizes feasible interceptor allocation and then minimizes total assignment cost. Compared with the original MinMax method, the proposed formulation provides an alternative objective that seeks to assign as many interceptors as possible across the affected targets and then minimize total bidding cost, while also handling the DES case in which later requests compete with already committed interceptors. Under the shared simulation communication assumptions, the results provide an algorithmic comparison between MinMax and global assignment. However, real deployment feasibility still requires communication-aware validation, including bid collection, final assignment communication, synchronization, and confirmation latency on target hardware.",
    ),
]


TABLE_EDITS = [
    TableEdit(1, 3, 0, "Table 2 algorithm label", "Global", "Global Assignment"),
    TableEdit(2, 3, 0, "Table 4 row 1 method label", "MCMF", "Global Assignment"),
    TableEdit(2, 4, 0, "Table 4 row 2 method label", "MCMF", "Global Assignment"),
]


def normalize(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def tokens(text: str) -> list[str]:
    return re.findall(r"\s+|[^\s]+", text)


def diff_markup(original: str, revised: str, side: str) -> str:
    a = tokens(original)
    b = tokens(revised)
    seq = SequenceMatcher(None, a, b)
    out: list[str] = []
    for tag, i1, i2, j1, j2 in seq.get_opcodes():
        if tag == "equal":
            segment = "".join(a[i1:i2])
            out.append(html.escape(segment))
        elif tag == "delete" and side == "old":
            segment = "".join(a[i1:i2])
            out.append(f'<span class="del">{html.escape(segment)}</span>')
        elif tag == "insert" and side == "new":
            segment = "".join(b[j1:j2])
            out.append(f'<span class="ins">{html.escape(segment)}</span>')
        elif tag == "replace":
            if side == "old":
                segment = "".join(a[i1:i2])
                out.append(f'<span class="del">{html.escape(segment)}</span>')
            else:
                segment = "".join(b[j1:j2])
                out.append(f'<span class="ins">{html.escape(segment)}</span>')
    return "".join(out)


def build_html(changes: list[dict[str, str]]) -> str:
    items = []
    for change in changes:
        items.append(
            f"""
            <section class="change-card">
              <h2>{html.escape(change["label"])}</h2>
              <div class="grid">
                <div>
                  <h3>Original</h3>
                  <p>{diff_markup(change["old"], change["new"], "old")}</p>
                </div>
                <div>
                  <h3>Revised</h3>
                  <p>{diff_markup(change["old"], change["new"], "new")}</p>
                </div>
              </div>
            </section>
            """
        )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>WSN Report Changes</title>
  <style>
    body {{
      font-family: Arial, sans-serif;
      margin: 32px;
      color: #1f2937;
      background: #f8fafc;
      line-height: 1.6;
    }}
    h1 {{
      margin-bottom: 8px;
    }}
    .meta {{
      margin-bottom: 24px;
      color: #475569;
    }}
    .change-card {{
      background: white;
      border: 1px solid #dbe4ee;
      border-radius: 12px;
      padding: 20px;
      margin-bottom: 18px;
      box-shadow: 0 4px 16px rgba(15, 23, 42, 0.05);
    }}
    .grid {{
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 20px;
    }}
    h2 {{
      font-size: 20px;
      margin-top: 0;
    }}
    h3 {{
      font-size: 15px;
      margin-bottom: 8px;
      color: #334155;
    }}
    p {{
      white-space: pre-wrap;
      margin: 0;
      padding: 14px;
      background: #f8fafc;
      border-radius: 8px;
      border: 1px solid #e2e8f0;
    }}
    .del {{
      background: #fee2e2;
      color: #991b1b;
      text-decoration: line-through;
      padding: 1px 2px;
      border-radius: 3px;
    }}
    .ins {{
      background: #fef3c7;
      color: #92400e;
      padding: 1px 2px;
      border-radius: 3px;
      font-weight: 600;
    }}
    @media (max-width: 900px) {{
      .grid {{
        grid-template-columns: 1fr;
      }}
    }}
  </style>
</head>
<body>
  <h1>Highlighted Report Revisions</h1>
  <div class="meta">
    <div>Source: {html.escape(str(SOURCE_DOC))}</div>
    <div>Revised copy: {html.escape(str(OUTPUT_DOC))}</div>
    <div>Changed items shown below: {len(changes)}</div>
  </div>
  {''.join(items)}
</body>
</html>
"""


def main() -> None:
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SOURCE_DOC, OUTPUT_DOC)

    original_doc = Document(str(SOURCE_DOC))
    revised_doc = Document(str(OUTPUT_DOC))

    changes: list[dict[str, str]] = []

    for edit in PARAGRAPH_EDITS:
        original_para = original_doc.paragraphs[edit.index]
        revised_para = revised_doc.paragraphs[edit.index]
        original_text = original_para.text.replace("\xa0", " ").strip()
        if edit.expected not in original_text:
            raise ValueError(
                f"Paragraph {edit.index} did not match expectation for {edit.label!r}. "
                f"Found: {original_text!r}"
            )
        changes.append({"label": edit.label, "old": original_text, "new": edit.new_text})
        set_paragraph_text(revised_para, edit.new_text)

    for edit in TABLE_EDITS:
        original_cell = original_doc.tables[edit.table_idx].cell(edit.row_idx, edit.col_idx)
        revised_cell = revised_doc.tables[edit.table_idx].cell(edit.row_idx, edit.col_idx)
        original_text = original_cell.text.replace("\xa0", " ").strip()
        if normalize(original_text) != normalize(edit.expected):
            raise ValueError(
                f"Table edit mismatch for {edit.label!r}. "
                f"Expected {edit.expected!r}, found {original_text!r}"
            )
        changes.append({"label": edit.label, "old": original_text, "new": edit.new_text})
        set_cell_text(revised_cell, edit.new_text)

    revised_doc.save(str(OUTPUT_DOC))

    revised_check = Document(str(OUTPUT_DOC))
    forbidden = [
        "mcmf",
        "min-cost max-flow",
        "flow network",
        "source/sink",
        "source sink",
        "edge capacities",
        "flow solver",
    ]
    body_text = []
    body_text.extend(p.text for p in revised_check.paragraphs)
    for table in revised_check.tables:
        for row in table.rows:
            for cell in row.cells:
                body_text.append(cell.text)
    joined = "\n".join(body_text).lower()
    hits = [term for term in forbidden if term in joined]
    if hits:
        raise ValueError(f"Forbidden terminology remained in revised document: {hits}")

    OUTPUT_HTML.write_text(build_html(changes), encoding="utf-8")

    print(f"Revised DOCX: {OUTPUT_DOC}")
    print(f"Changes HTML: {OUTPUT_HTML}")
    print(f"Changed items: {len(changes)}")


if __name__ == "__main__":
    main()
